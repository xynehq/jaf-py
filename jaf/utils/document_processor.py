"""
Document processing utilities for the JAF framework.

This module provides robust document content extraction with support for various
formats including PDF, Word documents, Excel spreadsheets, and more.
"""

import asyncio
import base64
import csv
import io
import json
import zipfile
from typing import Dict, Any, Optional, List

try:
    import aiofiles

    HAS_AIOFILES = True
except ImportError:
    HAS_AIOFILES = False

try:
    import httpx

    HAS_HTTPX = True
except ImportError:
    HAS_HTTPX = False
from pydantic import BaseModel

from ..core.types import Attachment

# Optional imports with graceful fallbacks
try:
    import PyPDF2

    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook

    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    import magic

    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

try:
    from PIL import Image

    HAS_PIL = True
except ImportError:
    HAS_PIL = False


# Constants
FETCH_TIMEOUT = 30.0
MAX_DOCUMENT_SIZE = 25 * 1024 * 1024  # 25MB
MAX_CSV_PREVIEW_ROWS = 10
MAX_EXCEL_SHEETS = 3
MAX_EXCEL_ROWS_PER_SHEET = 20


class DocumentProcessingError(Exception):
    """Exception raised when document processing fails."""

    def __init__(self, message: str, cause: Optional[Exception] = None):
        super().__init__(message)
        self.cause = cause


class NetworkError(Exception):
    """Exception raised when network operations fail."""

    def __init__(self, message: str, status_code: Optional[int] = None):
        super().__init__(message)
        self.status_code = status_code


class ProcessedDocument(BaseModel):
    """Result of document processing."""

    content: str
    metadata: Optional[Dict[str, Any]] = None


async def _fetch_url_content(url: str) -> tuple[bytes, Optional[str]]:
    """
    Fetch content from URL and return as bytes with content type.

    Args:
        url: URL to fetch

    Returns:
        Tuple of (content_bytes, content_type)

    Raises:
        NetworkError: If fetch fails
        DocumentProcessingError: If file is too large
    """
    if not HAS_HTTPX:
        raise DocumentProcessingError(
            "URL fetching not available. Install with: pip install 'jaf-py[attachments]'"
        )

    try:
        async with httpx.AsyncClient(timeout=FETCH_TIMEOUT) as client:
            # First check content length with a HEAD request if possible
            try:
                head_response = await client.head(
                    url,
                    headers={"User-Agent": "JAF-DocumentProcessor/1.0"},
                    timeout=FETCH_TIMEOUT / 2,  # Shorter timeout for HEAD request
                )
                head_response.raise_for_status()

                # Check Content-Length header if present
                content_length_str = head_response.headers.get("content-length")
                if content_length_str and content_length_str.isdigit():
                    content_length = int(content_length_str)
                    if content_length > MAX_DOCUMENT_SIZE:
                        size_mb = round(content_length / 1024 / 1024)
                        max_mb = round(MAX_DOCUMENT_SIZE / 1024 / 1024)
                        raise DocumentProcessingError(
                            f"File size ({size_mb}MB) exceeds maximum allowed size ({max_mb}MB)"
                        )
            except (httpx.HTTPStatusError, httpx.RequestError):
                # HEAD request failed, we'll check size during streaming
                pass

            # Stream the response to validate size as we download
            content_type = None
            accumulated_bytes = bytearray()
            async with client.stream(
                "GET", url, headers={"User-Agent": "JAF-DocumentProcessor/1.0"}
            ) as response:
                response.raise_for_status()
                content_type = response.headers.get("content-type")

                # Process the response in chunks
                async for chunk in response.aiter_bytes(chunk_size=8192):
                    accumulated_bytes.extend(chunk)
                    if len(accumulated_bytes) > MAX_DOCUMENT_SIZE:
                        size_mb = round(len(accumulated_bytes) / 1024 / 1024)
                        max_mb = round(MAX_DOCUMENT_SIZE / 1024 / 1024)
                        raise DocumentProcessingError(
                            f"File size ({size_mb}MB) exceeds maximum allowed size ({max_mb}MB)"
                        )

            return bytes(accumulated_bytes), content_type

    except httpx.HTTPStatusError as e:
        raise NetworkError(
            f"HTTP {e.response.status_code}: {e.response.reason_phrase}", e.response.status_code
        )
    except httpx.RequestError as e:
        raise NetworkError(f"Failed to fetch URL content: {e}", cause=e)
    except Exception as e:
        # Preserve system exceptions
        if isinstance(e, (KeyboardInterrupt, SystemExit, GeneratorExit, MemoryError)):
            raise
        raise NetworkError(f"Failed to fetch URL content: {e}", cause=e)


async def extract_document_content(attachment: Attachment) -> ProcessedDocument:
    """
    Extract text content from various document formats.

    Args:
        attachment: Attachment to process

    Returns:
        ProcessedDocument with extracted content

    Raises:
        DocumentProcessingError: If processing fails
    """
    # Get content as bytes
    if attachment.url and not attachment.data:
        content_bytes, detected_mime_type = await _fetch_url_content(attachment.url)
        mime_type = attachment.mime_type or detected_mime_type
    elif attachment.data:
        content_bytes = base64.b64decode(attachment.data)
        mime_type = attachment.mime_type
    else:
        raise DocumentProcessingError("No document data or URL provided")

    # Normalize MIME type
    mime_type = mime_type.lower() if mime_type else None

    # Process based on MIME type
    if mime_type == "application/pdf":
        return await _extract_pdf_content(content_bytes)
    elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        return _extract_docx_content(content_bytes)
    elif mime_type in [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ]:
        return _extract_excel_content(content_bytes)
    elif mime_type == "application/json":
        return _extract_json_content(content_bytes)
    elif mime_type == "application/zip":
        return _extract_zip_content(content_bytes)
    elif mime_type in ["text/plain", "text/csv"]:
        return _extract_text_content(content_bytes, mime_type)
    else:
        # Fallback: try to extract as text
        return _extract_text_content(content_bytes, "text/plain")


async def _extract_pdf_content(content_bytes: bytes) -> ProcessedDocument:
    """Extract content from PDF."""
    if not HAS_PDF:
        raise DocumentProcessingError(
            "PDF processing not available. Install with: pip install 'jaf-py[attachments]'"
        )

    try:
        # Run PDF processing in thread pool to avoid blocking
        def _process_pdf() -> ProcessedDocument:
            reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            text_parts = []

            for page in reader.pages:
                text_parts.append(page.extract_text())

            content = "\n".join(text_parts).strip()

            return ProcessedDocument(
                content=content,
                metadata={
                    "pages": len(reader.pages),
                    "info": dict(reader.metadata) if reader.metadata else None,
                },
            )

        return await asyncio.get_event_loop().run_in_executor(None, _process_pdf)

    except Exception as e:
        raise DocumentProcessingError(f"Failed to extract PDF content: {e}") from e


def _extract_text_content(content_bytes: bytes, mime_type: str) -> ProcessedDocument:
    """Extract content from text files."""
    try:
        content = content_bytes.decode("utf-8").strip()

        if mime_type == "text/csv":
            # Parse CSV to provide structured overview
            try:
                csv_reader = csv.DictReader(io.StringIO(content))
                rows = list(csv_reader)
                columns = csv_reader.fieldnames or []

                content_lines = content.split("\n")
                preview_lines = content_lines[:MAX_CSV_PREVIEW_ROWS]

                formatted_content = (
                    f"CSV File Content:\n"
                    f"Rows: {len(rows)}, Columns: {len(columns)}\n"
                    f"Columns: {', '.join(columns)}\n\n"
                    f"First few rows:\n{chr(10).join(preview_lines)}"
                )

                return ProcessedDocument(
                    content=formatted_content,
                    metadata={"rows": len(rows), "columns": len(columns), "fields": columns},
                )
            except Exception:
                # Fallback to raw text if CSV parsing fails
                pass

        return ProcessedDocument(content=content)

    except UnicodeDecodeError as e:
        raise DocumentProcessingError(f"Failed to decode text content: {e}") from e


def _extract_excel_content(content_bytes: bytes) -> ProcessedDocument:
    """Extract content from Excel files."""
    if not HAS_EXCEL:
        raise DocumentProcessingError(
            "Excel processing not available. Install with: pip install 'jaf-py[attachments]'"
        )

    try:
        workbook = load_workbook(io.BytesIO(content_bytes), read_only=True)
        sheet_names = workbook.sheetnames

        content_parts = [f"Excel File Content:\nSheets: {', '.join(sheet_names)}\n"]

        # Extract content from each sheet (limit to avoid overwhelming output)
        for i, sheet_name in enumerate(sheet_names):
            if i >= MAX_EXCEL_SHEETS:
                break

            worksheet = workbook[sheet_name]
            content_parts.append(f"\nSheet: {sheet_name}")

            # Extract up to MAX_EXCEL_ROWS_PER_SHEET rows
            rows_data = []
            for row_num, row in enumerate(worksheet.iter_rows(values_only=True), 1):
                if row_num > MAX_EXCEL_ROWS_PER_SHEET:
                    break
                # Convert row to strings, handling None values
                row_strings = [str(cell) if cell is not None else "" for cell in row]
                rows_data.append(",".join(row_strings))

            content_parts.append("\n".join(rows_data))

        content = "\n".join(content_parts).strip()

        return ProcessedDocument(content=content, metadata={"sheets": sheet_names})

    except Exception as e:
        raise DocumentProcessingError(f"Failed to extract Excel content: {e}") from e


def _extract_docx_content(content_bytes: bytes) -> ProcessedDocument:
    """Extract content from Word documents."""
    if not HAS_DOCX:
        raise DocumentProcessingError(
            "Word document processing not available. Install with: pip install 'jaf-py[attachments]'"
        )

    try:
        document = Document(io.BytesIO(content_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        content = "\n".join(paragraphs).strip()

        return ProcessedDocument(content=content, metadata={"paragraphs": len(paragraphs)})

    except Exception as e:
        raise DocumentProcessingError(f"Failed to extract DOCX content: {e}") from e


def _extract_json_content(content_bytes: bytes) -> ProcessedDocument:
    """Extract content from JSON files."""
    try:
        json_str = content_bytes.decode("utf-8")
        json_obj = json.loads(json_str)

        # Pretty print JSON with some metadata
        formatted_content = f"JSON File Content:\n{json.dumps(json_obj, indent=2)}"

        metadata = {"type": "array" if isinstance(json_obj, list) else type(json_obj).__name__}

        if isinstance(json_obj, dict):
            metadata["keys"] = list(json_obj.keys())
        elif isinstance(json_obj, list):
            metadata["length"] = len(json_obj)

        return ProcessedDocument(content=formatted_content, metadata=metadata)

    except (UnicodeDecodeError, json.JSONDecodeError):
        # Fallback to raw text if JSON parsing fails
        if isinstance(content_bytes, bytes):
            # If input is bytes, decode with error handling
            fallback_content = content_bytes.decode("utf-8", errors="replace").strip()
        else:
            # If input is already a string (from a previous decode attempt)
            fallback_content = json_str.strip() if isinstance(json_str, str) else str(content_bytes)

        return ProcessedDocument(content=fallback_content)


def _extract_zip_content(content_bytes: bytes) -> ProcessedDocument:
    """Extract file listing from ZIP archives."""
    try:
        with zipfile.ZipFile(io.BytesIO(content_bytes), "r") as zip_file:
            files = zip_file.namelist()

            content_parts = ["ZIP File Contents:\n"]
            safe_files = []

            # Create virtual root for path safety checks
            from pathlib import Path
            import os

            virtual_root = Path(
                "/safe_extract_dir"
            )  # Virtual root never actually used for extraction

            for file_name in files:
                # Skip empty entries
                if not file_name:
                    continue

                # Basic security checks
                if (
                    file_name.startswith("/")  # Absolute path
                    or file_name.startswith("\\")  # Windows absolute path
                    or file_name.startswith("..")  # Parent directory traversal
                    or ".." in file_name.split("/")  # Parent directory traversal
                    or ".." in file_name.split("\\")  # Windows traversal
                    or ":" in file_name  # Windows drive letter
                    or "\0" in file_name
                ):  # Null byte
                    # Skip unsafe entries
                    content_parts.append(f"WARNING: Skipped suspicious path: {file_name[:50]}...")
                    continue

                # Normalize path for additional safety check
                try:
                    # Create safe path relative to virtual root
                    norm_path = os.path.normpath(file_name)
                    if norm_path.startswith(".."):
                        # Skip unsafe entries that normalize to traversal
                        content_parts.append(
                            f"WARNING: Skipped path traversal attempt: {file_name[:50]}..."
                        )
                        continue

                    # Check if path would escape the virtual root
                    test_path = virtual_root.joinpath(norm_path).resolve()
                    if not str(test_path).startswith(str(virtual_root)):
                        # Skip unsafe entries that would escape extraction root
                        content_parts.append(
                            f"WARNING: Skipped path traversal attempt: {file_name[:50]}..."
                        )
                        continue

                    # Passed all security checks, add to safe file list
                    safe_files.append(file_name)

                    # Get file info for display
                    if file_name.endswith("/"):
                        content_parts.append(f"DIR: {file_name}")
                    else:
                        try:
                            file_info = zip_file.getinfo(file_name)
                            size = file_info.file_size
                            content_parts.append(f"FILE: {file_name} ({size} bytes)")
                        except KeyError:
                            content_parts.append(f"FILE: {file_name}")
                except Exception:
                    # Skip any entry that causes normalization errors
                    content_parts.append(f"WARNING: Skipped invalid path: {file_name[:50]}...")
                    continue

            content = "\n".join(content_parts).strip()

            return ProcessedDocument(
                content=content, metadata={"files": safe_files, "total_files": len(safe_files)}
            )

    except Exception as e:
        raise DocumentProcessingError(f"Failed to process ZIP file: {e}") from e


def is_document_supported(mime_type: Optional[str]) -> bool:
    """
    Check if a MIME type is supported for content extraction.

    Args:
        mime_type: MIME type to check

    Returns:
        True if supported, False otherwise
    """
    if not mime_type:
        return False

    supported_types = [
        "application/pdf",
        "text/plain",
        "text/csv",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/json",
        "application/zip",
    ]

    return mime_type.lower() in supported_types


def get_document_description(mime_type: Optional[str]) -> str:
    """
    Get a human-readable description of what content will be extracted.

    Args:
        mime_type: MIME type to describe

    Returns:
        Human-readable description
    """
    if not mime_type:
        return "document content"

    descriptions = {
        "application/pdf": "PDF text content",
        "text/plain": "plain text content",
        "text/csv": "CSV data structure and sample rows",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "Excel spreadsheet data",
        "application/vnd.ms-excel": "Excel spreadsheet data",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "Word document text content",
        "application/json": "JSON data structure",
        "application/zip": "ZIP file listing",
    }

    return descriptions.get(mime_type.lower(), "document content")


def get_missing_dependencies() -> List[str]:
    """
    Get list of missing optional dependencies for document processing.

    Returns:
        List of missing dependency names
    """
    missing = []

    if not HAS_PDF:
        missing.append("PyPDF2 (for PDF processing)")
    if not HAS_DOCX:
        missing.append("python-docx (for Word document processing)")
    if not HAS_EXCEL:
        missing.append("openpyxl (for Excel processing)")
    if not HAS_PIL:
        missing.append("Pillow (for image processing)")
    if not HAS_MAGIC:
        missing.append("python-magic (for MIME type detection)")
    if not HAS_HTTPX:
        missing.append("httpx (for URL fetching)")
    if not HAS_AIOFILES:
        missing.append("aiofiles (for async file operations)")

    return missing


def check_dependencies() -> Dict[str, bool]:
    """
    Check availability of optional dependencies.

    Returns:
        Dictionary mapping dependency names to availability
    """
    return {
        "pdf": HAS_PDF,
        "docx": HAS_DOCX,
        "excel": HAS_EXCEL,
        "image": HAS_PIL,
        "magic": HAS_MAGIC,
        "httpx": HAS_HTTPX,
        "aiofiles": HAS_AIOFILES,
    }
